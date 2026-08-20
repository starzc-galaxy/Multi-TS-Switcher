"""生成《Multi-TS Switcher 使用说明书》docx（python-docx）。"""

import sys
from datetime import date, datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT2 = RGBColor(0x25, 0x63, 0xEB)
GRAY = RGBColor(0x6B, 0x72, 0x80)


def set_east_asia(style_or_run, name: str = "微软雅黑") -> None:
    rpr = style_or_run.font.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(10.5)
    set_east_asia(normal)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(4)

    for name, size, before, after in (
        ("Heading 1", 16, 18, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11.5, 8, 4),
    ):
        st = doc.styles[name]
        st.font.name = "微软雅黑"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = ACCENT
        set_east_asia(st)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def add_page_number_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("第 ")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    fld = p.add_run()
    fld.font.size = Pt(9)
    fld.font.color.rgb = GRAY
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    fld._r.append(fc1)
    fld._r.append(instr)
    fld._r.append(fc2)
    tail = p.add_run(" 页")
    tail.font.size = Pt(9)
    tail.font.color.rgb = GRAY


def shade_cell(cell, hex_color: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def set_cell(cell, text: str, bold: bool = False, color: RGBColor | None = None,
             size: float = 10.5, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "微软雅黑"
    set_east_asia(run)
    if color is not None:
        run.font.color.rgb = color


def add_table(doc: Document, headers: list[str], rows: list[list],
              widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for j, w in enumerate(widths_cm):
        for cell in table.columns[j].cells:
            cell.width = Cm(w)
    for j, h in enumerate(headers):
        set_cell(table.rows[0].cells[j], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                 center=True)
        shade_cell(table.rows[0].cells[j], "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            set_cell(cells[j], val, center=(j == 0 or len(str(val)) <= 12))
    # 表头跨页重复
    trpr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


def numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(2)
    return p


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def build(out_path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    setup_styles(doc)
    add_page_number_footer(doc)

    # ---------- 封面 ----------
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Multi-TS Switcher")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = ACCENT2
    r.font.name = "微软雅黑"
    set_east_asia(r)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = st.add_run("多路 UDP-TS 轮询切换转发系统 · 使用说明书")
    r2.font.size = Pt(16)
    r2.font.color.rgb = GRAY
    r2.font.name = "微软雅黑"
    set_east_asia(r2)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from app.version import APP_VERSION

    rm = meta.add_run(f"版本 V{APP_VERSION}　|　{date.today().isoformat()}")
    rm.font.size = Pt(11)
    rm.font.color.rgb = GRAY
    doc.add_page_break()

    # ---------- 目录 ----------
    h1(doc, "目录")
    toc = [
        "1  系统简介",
        "2  快速开始",
        "3  软件授权",
        "4  主程序界面",
        "5  组配置",
        "6  运行与切换机制",
        "7  测试工具",
        "8  授权生成器（管理员）",
        "9  日志与故障排查",
        "10 安全与维护提示",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
    doc.add_page_break()

    # ---------- 1 系统简介 ----------
    h1(doc, "1　系统简介")
    doc.add_paragraph(
        "Multi-TS Switcher 用于多路 UDP-TS 信号按组轮询切换后输出单路 UDP-TS 的播出调度场景。"
        "软件不转码，只做 TS 包级切换转发；切换时对 PCR/PTS 时间戳进行重基准并等待关键帧，"
        "尽可能避免下游画面卡顿或黑屏。"
    )
    h2(doc, "1.1 主要功能")
    bullet(doc, "最多 9 组并行，每组最多 9 个输入源，外加 1 个垫片源（本地 TS 文件兜底）。")
    bullet(doc, "输入/输出均支持单播与组播，可按网卡绑定收发。")
    bullet(doc, "每组按配置间隔轮询切换，自动跳过异常源；全部异常自动切垫片。")
    bullet(doc, "关键帧对齐切换 + PCR/PTS/DTS 重基准 + discontinuity 标记 + CC 重置。")
    bullet(doc, "每路源实时监测：无数据超时、CC 错误、PCR 抖动、码率、收包数。")
    bullet(doc, "监控墙响应式布局（1–4 列自适应），F11 全屏只显示视频墙。")
    bullet(doc, "卡片画面右上角显示组信息、正常/异常源数量、当前源名称。")
    bullet(doc, "JSON 配置热生效；按模块本地日志；异常退出原因自动记录。")
    bullet(doc, "离线授权：机器指纹 + 签名授权文件，授权组数 1–9，可设有效期（天/永久）。")
    h2(doc, "1.2 交付物")
    add_table(
        doc,
        ["程序", "用途", "说明"],
        [
            ["MultiTS_Switcher.exe", "主程序", "多组轮询切换转发与监控界面"],
            ["TestTool.exe", "测试工具", "发送测试源、接收验证主程序输出"],
            ["LicenseGenerator.exe", "授权生成器", "管理员签发授权文件（勿发给客户）"],
            ["使用说明书.docx", "本文档", "安装、配置、操作与排障说明"],
        ],
        [4.2, 4.2, 8.2],
    )

    # ---------- 2 快速开始 ----------
    h1(doc, "2　快速开始")
    numbered(doc, "双击 MultiTS_Switcher.exe 启动；首次运行请先导入授权文件（见第 3 章）。")
    numbered(doc, "在右侧“组设置”中配置输出地址（如 230.1.1.1:7000）和输入源地址。")
    numbered(doc, "点击左上角“启动全部”开关，引擎开始接收并转发。")
    numbered(doc, "卡片画面显示当前源预览；右上角浮层显示组信息、正常/异常数量和当前源。")
    numbered(doc, "如需验证，打开 TestTool.exe 的“接收验证”，填入主程序输出地址并开始接收。")

    # ---------- 3 软件授权 ----------
    h1(doc, "3　软件授权")
    h2(doc, "3.1 授权文件")
    doc.add_paragraph(
        "软件采用离线授权：授权文件（.lic）绑定本机机器码，由管理员用授权生成器签发。"
        "未授权时程序只显示授权窗口，不能运行。"
    )
    h2(doc, "3.2 导入授权")
    numbered(doc, "在授权窗口查看本机机器码，点“复制”按钮复制完整机器码并发给管理员。")
    numbered(doc, "管理员用授权生成器签发 .lic 文件后发回。")
    numbered(doc, "点击“导入授权文件…”，选择 .lic 文件，导入成功后自动进入主界面。")
    h2(doc, "3.3 查看与重新导入")
    doc.add_paragraph(
        "工具栏“授权”可查看当前机器码、授权组数和有效期；需要升级/调整组数或续期时，"
        "点击“导入新授权…”重新导入即可立即生效，超出新组数的引擎会自动停止。"
    )
    h2(doc, "3.4 有效期")
    doc.add_paragraph(
        "授权文件可设为永久或 N 天有效。到期后程序会提示“授权已过期”并回到授权窗口，"
        "重新导入有效授权即可恢复。"
    )

    # ---------- 4 主程序界面 ----------
    h1(doc, "4　主程序界面")
    h2(doc, "4.1 界面布局")
    bullet(doc, "上方工具栏：启动/停止开关、设置、日志、授权、全屏。")
    bullet(doc, "左侧监控墙：所有组卡片按窗口宽度自动排列 1–4 列。")
    bullet(doc, "右侧组设置：查看/编辑组参数、输入源与实时状态。")
    h2(doc, "4.2 卡片信息浮层")
    doc.add_paragraph(
        "每张卡片右上角直接叠在画面上：组名、输出地址、正常/异常源数量、当前源名称。"
        "画面为当前输出源的小预览，随轮询切换自动更新。"
    )
    h2(doc, "4.3 全屏模式")
    doc.add_paragraph("按 F11 或点击“全屏”只显示监控墙，适合值班大屏；再次按 F11 退出。")

    # ---------- 5 组配置 ----------
    h1(doc, "5　组配置")
    h2(doc, "5.1 组参数")
    add_table(
        doc,
        ["参数", "说明"],
        [
            ["组名 / 组备注", "用于识别；备注会显示在组设置和卡片浮层中"],
            ["轮询间隔", "每组统一轮询周期（秒），到点切换到下一个正常源"],
            ["输出地址 / 端口", "本组单路 UDP-TS 输出目标，可选单播或组播"],
            ["绑定网卡", "下拉选择本机网卡，用于组播加入与发送；默认自动"],
            ["垫片文件", "所有源异常时循环发送的本地 TS 文件，默认黑场+静音"],
            ["启用本组", "关闭后本组不参与运行"],
        ],
        [4.0, 12.6],
    )
    h2(doc, "5.2 输入源与实时状态")
    doc.add_paragraph(
        "输入源表格一行一个源，前几列（名称/地址/端口/备注）可编辑，后几列为实时状态："
        "状态灯、码率、收包数、CC 错误、断流。新加入的源默认启用、默认组播。"
    )
    h2(doc, "5.3 保存与热生效")
    doc.add_paragraph(
        "修改后点“保存配置”即写入 JSON 并热推给对应引擎进程，无需重启；"
        "手动编辑 config/groups.json 也会被读取（重启程序或重新保存后生效）。"
    )

    # ---------- 6 运行与切换机制 ----------
    h1(doc, "6　运行与切换机制")
    h2(doc, "6.1 轮询策略")
    bullet(doc, "到轮询点后从健康源列表中选择下一个源，异常源自动跳过。")
    bullet(doc, "正常切换等待新源的关键帧（可随机接入点）后再切，等待期间继续输出当前源。")
    bullet(doc, "全部源异常时立即切到垫片源；源恢复后重新参与轮询。")
    h2(doc, "6.2 时间戳处理")
    doc.add_paragraph(
        "切换瞬间记录输出侧时间基准，计算偏移后对源的全部 PCR/PTS/DTS 持续重基准，"
        "保证输出时间戳单调连续；切换首包写入 discontinuity 标记并重置 CC 计数，"
        "避免下游解码器因时间戳回跳而卡顿。"
    )
    h2(doc, "6.3 异常与兜底")
    doc.add_paragraph(
        "每路源按“无数据超时”判定断流；异常源在界面统计中显示为异常。"
        "垫片文件缺失时该组不再自动兜底，请确认 assets/filler.ts 存在。"
    )

    # ---------- 7 测试工具 ----------
    h1(doc, "7　测试工具")
    doc.add_paragraph(
        "TestTool.exe 用于联调：一端发送测试源，一端接收验证主程序输出。"
    )
    h2(doc, "7.1 发送测试源")
    bullet(doc, "选择路数（1–9）、起始地址与端口，默认 229.1.1.1:7000 组播。")
    bullet(doc, "每路源颜色不同且带移动竖条，便于肉眼确认切换到了哪一路。")
    bullet(doc, "测试源文件首次使用自动生成到程序目录 test_sources/。")
    h2(doc, "7.2 接收验证")
    bullet(doc, "填入主程序输出地址（如 230.1.1.1:7000），点击“开始接收”。")
    bullet(doc, "实时解码显示画面，并统计收包数、码率、CC 错误。")
    bullet(doc, "主程序切换源时，接收画面会跟随切换后的画面更新。")

    # ---------- 8 授权生成器 ----------
    h1(doc, "8　授权生成器（管理员）")
    numbered(doc, "打开 LicenseGenerator.exe，选择私钥文件（首次使用需生成密钥对）。")
    numbered(doc, "粘贴客户提供的机器码（64 位十六进制），选择授权组数（1–9）。")
    numbered(doc, "选择有效期（永久或按天数），点“生成授权”，得到 .lic 文件。")
    numbered(doc, "把 .lic 文件发给客户导入。")
    doc.add_paragraph(
        "重要：本交付包内的授权生成器已自带私钥文件（dev_private_key.pem），可直接签发授权。"
        "私钥等于签发权，请把整个授权工具视为机密交付物，只交给授权管理员保管，切勿随软件随意转传；"
        "私钥泄露等于任何人都能签发授权。"
    )

    # ---------- 9 日志与故障排查 ----------
    h1(doc, "9　日志与故障排查")
    h2(doc, "9.1 日志文件")
    doc.add_paragraph(
        "运行日志位于程序目录 logs/：ui.log（界面）、engine.log（引擎）、"
        "switch.log（切换记录）、error.log（错误与崩溃原因）。异常退出原因会自动写入 error.log。"
    )
    h2(doc, "9.2 常见问题")
    add_table(
        doc,
        ["现象", "原因与处理"],
        [
            ["收不到组播源", "检查防火墙是否允许 UDP 入站；多网卡环境在“绑定网卡”选择正确的网卡"],
            ["某路源一直异常", "确认该组播地址有实际数据；组与组之间不要复用同一个组播地址"],
            ["画面卡住不更新", "多为源端中断或 PCR 异常；查看 engine.log 与 switch.log 定位"],
            ["垫片不生效", "确认垫片文件路径存在且为有效 TS（默认 assets/filler.ts）"],
            ["授权提示过期/无效", "重新导入有效授权；确认机器码与本机一致"],
            ["程序启动停在授权窗口", "正常流程，导入授权后进入主界面"],
        ],
        [4.6, 12.0],
    )

    # ---------- 10 安全与维护提示 ----------
    h1(doc, "10　安全与维护提示")
    bullet(doc, "私钥文件（dev_private_key.pem）随授权工具进包；整个授权工具属于机密交付物，只交给授权管理员，切勿外传。")
    bullet(doc, "生产环境建议为每组使用不同的组播地址，避免同端口组播互相干扰。")
    bullet(doc, "多网卡机房建议显式选择绑定网卡，避免组播走到错误网卡。")
    bullet(doc, "升级程序前先停止全部引擎并关闭旧程序，再替换 exe。")
    bullet(doc, "定期查看 logs/error.log，异常退出记录有助于快速定位问题。")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out_path))
    except PermissionError:
        alt = out_path.with_name(
            f"使用说明书_{datetime.now():%Y%m%d_%H%M%S}.docx"
        )
        doc.save(str(alt))
        out_path = alt
        print(f"目标文件被占用，已改存: {alt}")
    print(f"manual written: {out_path} ({out_path.stat().st_size} bytes)")
    return out_path


def verify(out_path: Path) -> None:
    d = Document(str(out_path))
    headings = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
    print("headings:", len(headings))
    for h in headings:
        print(" -", h)
    print("tables:", len(d.tables))
    assert headings
    assert any("系统简介" in h for h in headings)
    assert any("故障排查" in h for h in headings)
    assert len(d.tables) >= 3
    print("verify ok")


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("交付") / "使用说明书.docx"
    final = build(out)
    verify(final)
